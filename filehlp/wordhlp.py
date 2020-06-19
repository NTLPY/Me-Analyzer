"""
    Word Helper

    This header is to extract the infomation included in word files.

    Author:   NTLPY
    Creation: 0x01d6456349a27230 (UTC)

"""

# import modules
import wcuthlp;
import docx as _docx;
import win32com.client as _w32c_client;

# kernel object
#word_client = _w32c_client.Dispatch("wps.Application");

# Initialization
temp_file_name = "temp"

# transform a doc into a docx
def DOC_to_DOCX(source : str, destination : str):
    file = word_client.Documents.Open(source);
    file.SaveAs(destination, 12);
    return destination;

# extract from docx
def get_from_docx(filename : str, words : dict):
    file = _docx.Document(filename);
    for p in file.paragraphs:
        for r in p.runs:
            wcuthlp.cut_words(r.text, words);

    return;

# extract from doc
def get_from_doc(filename : str, words : dict):
    DOC_to_DOCX(filename, "temp");

    file = _docx.Document("temp");
    for p in file.paragraphs:
        for r in p.runs:
            wcuthlp.cut_words(r.text, dict);

    return;
